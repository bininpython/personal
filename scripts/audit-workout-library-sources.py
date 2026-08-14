#!/usr/bin/env python3
"""Audita PDFs e DOCX usados como fontes da biblioteca de treinos."""

from __future__ import annotations

import argparse
import hashlib
import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SOURCE_EXTENSIONS = {".pdf", ".docx"}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def extract_pdf(path: Path) -> dict[str, object]:
    reader = PdfReader(path)
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        pages.append({
            "page": index,
            "text": (page.extract_text() or "").strip(),
        })
    return {
        "kind": "pdf",
        "pageCount": len(reader.pages),
        "pages": pages,
    }


def extract_docx(path: Path) -> dict[str, object]:
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"table": table_index, "rows": rows})

    inline_shapes = len(document.inline_shapes)
    embedded_images = 0
    for relationship in document.part.rels.values():
        if "image" in relationship.reltype:
            embedded_images += 1

    return {
        "kind": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
        "inlineShapeCount": inline_shapes,
        "embeddedImageCount": embedded_images,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source_root", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    source_root = args.source_root.resolve()
    output = args.output.resolve()
    files = sorted(
        (
            path for path in source_root.rglob("*")
            if path.is_file()
            and path.suffix.lower() in SOURCE_EXTENSIONS
            and "fitcontrol-pro" not in path.parts
        ),
        key=lambda item: str(item.relative_to(source_root)).casefold(),
    )

    groups: dict[str, list[Path]] = defaultdict(list)
    for path in files:
        groups[sha256(path)].append(path)

    records = []
    for file_hash, duplicates in sorted(groups.items(), key=lambda item: str(item[1][0]).casefold()):
        canonical = min(duplicates, key=lambda item: (len(item.parts), len(str(item)), str(item).casefold()))
        extracted = extract_pdf(canonical) if canonical.suffix.lower() == ".pdf" else extract_docx(canonical)
        records.append({
            "sha256": file_hash,
            "canonicalPath": str(canonical.relative_to(source_root)).replace("\\", "/"),
            "duplicatePaths": [
                str(path.relative_to(source_root)).replace("\\", "/")
                for path in duplicates
            ],
            **extracted,
        })

    payload = {
        "sourceRoot": str(source_root),
        "fileCount": len(files),
        "uniqueFileCount": len(records),
        "sources": records,
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8", newline="\n")
    print(json.dumps({
        "fileCount": len(files),
        "uniqueFileCount": len(records),
        "output": str(output),
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
